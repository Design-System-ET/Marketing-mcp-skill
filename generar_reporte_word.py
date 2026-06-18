#!/usr/bin/env python3
"""
generar_reporte_word.py — Genera un documento Word profesional (.docx) con
portada, numeración de páginas, pie de página y todo el contenido del
onboarding estructurado con títulos, tablas y secciones.

Modo de uso:
  python generar_reporte_word.py --cliente MiPrueba
  python generar_reporte_word.py --cliente MiPrueba --output "MiInforme.docx"

Requiere: python-docx, beautifulsoup4
"""

import os
import re
import sys
import logging
import argparse
from datetime import datetime
from glob import glob
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _aplicar_bordes_tabla(table, size: str = '4', color: str = 'D0D5DD', clean_first: bool = False):
    """Aplica bordes consistentes a una tabla de python-docx.

    Args:
        table: Objeto Table de python-docx.
        size: Grosor del borde en units (default '4').
        color: Color hex del borde (default 'D0D5DD').
        clean_first: Si True, elimina bordes existentes antes de aplicar.
    """
    tbl_pr = table._tbl.tblPr
    if clean_first:
        for old in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(old)
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), color)
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


# ═══════════════════════════════════════════════════════════════
# CONFIGURACIÓN DE ESTILOS — Paleta premium
# ═══════════════════════════════════════════════════════════════

COLOR_PRIMARY      = RGBColor(0x00, 0x4D, 0x9E)   # Azul corporativo profundo
COLOR_PRIMARY_LIGHT= RGBColor(0x00, 0x7B, 0xFF)   # Azul brillante
COLOR_PRIMARY_DARK = RGBColor(0x00, 0x2B, 0x5C)   # Azul oscuro
COLOR_ACCENT       = RGBColor(0xD4, 0xA0, 0x30)   # Dorado/copper
COLOR_DARK         = RGBColor(0x1A, 0x1A, 0x2E)   # Casi negro
COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GRAY         = RGBColor(0x88, 0x88, 0x88)
COLOR_GRAY_LIGHT   = RGBColor(0xAA, 0xAA, 0xAA)
COLOR_BG_LIGHT     = RGBColor(0xF5, 0xF7, 0xFA)   # Fondo muy claro
COLOR_BG_ALT       = RGBColor(0xE8, 0xEE, 0xF4)   # Fondo alternativo
COLOR_TEXT         = RGBColor(0x2D, 0x2D, 0x2D)   # Texto principal
COLOR_TEXT_MUTED   = RGBColor(0x66, 0x66, 0x66)   # Texto secundario

# Colores FODA
COLOR_F_FORTALEZA   = RGBColor(0x05, 0x96, 0x69)
COLOR_F_DEBILIDAD   = RGBColor(0xDC, 0x26, 0x26)
COLOR_F_OPORTUNIDAD = RGBColor(0x25, 0x63, 0xEB)
COLOR_F_AMENAZA     = RGBColor(0xD9, 0x77, 0x06)

# HEX para XML
HEX_TABLE_HEADER = "002B5C"
HEX_ACCENT       = "D4A030"

FONT_TITLE = 'Calibri Light'
FONT_BODY  = 'Calibri'

RUTA_CLIENTES = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Clientes')


# ═══════════════════════════════════════════════════════════════
# FUNCIONES AUXILIARES XML
# ═══════════════════════════════════════════════════════════════

def _add_field(paragraph, field_code: str):
    """Agrega un campo Word (PAGE, NUMPAGES, DATE) a un párrafo."""
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    run2 = paragraph.add_run()
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = field_code; run2._r.append(i)
    run3 = paragraph.add_run()
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); run3._r.append(f2)


def _set_cell_shading(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)


def _set_paragraph_shading(paragraph, hex_color: str):
    """Aplica color de fondo a un párrafo."""
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color); s.set(qn('w:val'), 'clear')
    paragraph._p.get_or_add_pPr().append(s)


def _set_paragraph_border(paragraph, left_color: str | None = None, left_size: int | None = None,
                          top_color: str | None = None, bottom_color: str | None = None):
    """Agrega bordes a un párrafo. left_color+left_size para barra vertical."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    if left_color and left_size:
        el = OxmlElement('w:left')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(left_size))
        el.set(qn('w:space'), '8'); el.set(qn('w:color'), left_color)
        borders.append(el)
    if top_color:
        el = OxmlElement('w:top')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '4'); el.set(qn('w:color'), top_color)
        borders.append(el)
    if bottom_color:
        el = OxmlElement('w:bottom')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '4'); el.set(qn('w:color'), bottom_color)
        borders.append(el)
    pPr.append(borders)


def _set_cell_vertical_alignment(cell, align: str = 'center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val = OxmlElement('w:vAlign')
    val.set(qn('w:val'), align)
    tcPr.append(val)


# ═══════════════════════════════════════════════════════════════
# FORMATEO DE CELDAS Y PÁRRAFOS
# ═══════════════════════════════════════════════════════════════

def _make_cell(cell, text: str, bold: bool = False, color=None, size: int = 10, align=None,
               font_name: str = FONT_BODY, italic: bool = False):
    """Formatea una celda limpiándola primero."""
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color


def _make_paragraph(doc, text: str, bold: bool = False, color=None, size: int = 11, align=None,
                    font_name: str = FONT_BODY, space_before: int = 0, space_after: int = 6, italic: bool = False):
    """Agrega un párrafo formateado."""
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    return p


def _crear_estilos(doc):
    """Configura los estilos base del documento."""
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    style.font.color.rgb = COLOR_TEXT

    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT_TITLE
        if level == 1:
            hs.font.size = Pt(20)
            hs.font.color.rgb = COLOR_PRIMARY_DARK
        elif level == 2:
            hs.font.size = Pt(15)
            hs.font.color.rgb = COLOR_PRIMARY
        elif level == 3:
            hs.font.size = Pt(13)
            hs.font.color.rgb = COLOR_PRIMARY_LIGHT


# ═══════════════════════════════════════════════════════════════
# HEADER DE PÁGINA
# ═══════════════════════════════════════════════════════════════

def _add_header_footer(doc, cliente: str):
    """Agrega header y footer a todas las secciones (menos portada)."""
    for i, section in enumerate(doc.sections):
        # ── Header ──
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(0)

        run = hp.add_run('Global Infinity Marketing')
        run.font.size = Pt(8); run.font.name = FONT_TITLE
        run.font.color.rgb = COLOR_GRAY; run.bold = True

        run = hp.add_run(f'  |  {cliente}')
        run.font.size = Pt(8); run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_GRAY_LIGHT

        # Línea inferior del header
        hp_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D4A030')
        hp_bdr.append(bottom)
        hp._p.get_or_add_pPr().append(hp_bdr)

        # ── Footer ──
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs: p.clear()
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after  = Pt(0)

        run = fp.add_run(f'{cliente}  |  Página ')
        run.font.size = Pt(8); run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_GRAY_LIGHT

        _add_field(fp, ' PAGE ')

        run = fp.add_run(f'  |  {datetime.now().strftime("%d/%m/%Y")}')
        run.font.size = Pt(8); run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_GRAY_LIGHT

        # Línea footer
        fp_bdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1'); top.set(qn('w:color'), 'CCCCCC')
        fp_bdr.append(top)
        fp._p.get_or_add_pPr().append(fp_bdr)


# ═══════════════════════════════════════════════════════════════
# PORTADA — Diseño premium
# ═══════════════════════════════════════════════════════════════

def _add_cover(doc, cliente: str, analista: str, fecha: str, total_pasos: int):
    """Portada profesional con diseño premium."""
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Barra superior decorativa ──
    p = doc.add_paragraph()
    _set_paragraph_border(p, top_color=HEX_ACCENT)
    p.paragraph_format.space_after = Pt(0)

    # ── Espaciado ──
    for _ in range(5):
        doc.add_paragraph('')

    # ── Línea dorada ──
    _make_paragraph(doc, '━' * 55, color=COLOR_ACCENT, size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── Título principal ──
    _make_paragraph(doc, 'INFORME DE ONBOARDING', bold=True,
                    color=COLOR_PRIMARY_DARK, size=34,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
                    font_name=FONT_TITLE)

    # ── Subtítulo ──
    _make_paragraph(doc, 'Diagnóstico Estratégico de Marketing Digital',
                    color=COLOR_ACCENT, size=16,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
                    font_name=FONT_TITLE, italic=True)

    # ── Línea dorada ──
    _make_paragraph(doc, '━' * 55, color=COLOR_ACCENT, size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # ── Caja de datos del informe ──
    # Crear tabla invisible de 1 celda para el fondo
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _set_cell_shading(cell, "F0F3F8")

    datos = [
        ('Cliente', cliente),
        ('Analista', analista),
        ('Fecha', fecha),
        ('Pasos completados', str(total_pasos)),
    ]
    # Primera fila
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)

    for idx, (label, value) in enumerate(datos):
        r = p.add_run(f'{label}: ')
        r.bold = True; r.font.size = Pt(11); r.font.name = FONT_BODY
        r.font.color.rgb = COLOR_PRIMARY_DARK
        r = p.add_run(value)
        r.font.size = Pt(11); r.font.name = FONT_BODY
        r.font.color.rgb = COLOR_TEXT_MUTED
        if idx < len(datos) - 1:
            r = p.add_run('    │    ')
            r.font.size = Pt(11); r.font.color.rgb = COLOR_ACCENT
            r.font.name = FONT_BODY

    # Segunda fila: tipo de documento
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(16)
    r = p2.add_run('Documento Word Profesional')
    r.font.size = Pt(9); r.font.name = FONT_BODY
    r.font.color.rgb = COLOR_GRAY; r.italic = True

    doc.add_paragraph('')

    # ── Línea dorada ──
    _make_paragraph(doc, '━' * 55, color=COLOR_ACCENT, size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ── Marca ──
    _make_paragraph(doc, 'Global Infinity Marketing', bold=True,
                    color=COLOR_PRIMARY, size=16,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
                    font_name=FONT_TITLE)

    _make_paragraph(doc, 'by Design System - Claudio Silveira',
                    color=COLOR_GRAY, size=9,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
                    font_name=FONT_BODY)

    # ── Barra inferior decorativa ──
    p_end = doc.add_paragraph()
    _set_paragraph_border(p_end, bottom_color=HEX_ACCENT)
    p_end.paragraph_format.space_before = Pt(0)

    # Salto de página
    doc.add_page_break()

    # Ahora agregar header/footer a la sección (después de la portada)
    # Esto se hace al final en la función principal


# ═══════════════════════════════════════════════════════════════
# SECCIÓN DE REPORTE — Con barra lateral decorativa
# ═══════════════════════════════════════════════════════════════

def _insert_seccion(doc, soup, paso: str, titulo: str, fase: str, cliente: str):
    """Agrega una sección completa con decoración lateral."""

    # ── Badge de PASO ──
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(4)
    p_badge.paragraph_format.space_after  = Pt(0)
    _set_paragraph_shading(p_badge, "002B5C")
    run = p_badge.add_run(f'  PASO {paso}  ')
    run.font.size = Pt(9); run.font.name = FONT_BODY
    run.font.color.rgb = COLOR_WHITE; run.bold = True

    # ── Título con barra lateral azul ──
    h1 = doc.add_heading(titulo, level=1)
    _set_paragraph_border(h1, left_color="D4A030", left_size=36)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after  = Pt(2)

    # ── Metadata en línea ──
    meta_items = soup.select('.header-meta-item')
    meta_data = {}
    for item in meta_items:
        strong = item.find('strong')
        if strong:
            label = strong.get_text(strip=True).replace(':', '')
            value = item.get_text(strip=True).replace(f'{label}:', '', 1).strip()
            meta_data[label] = value

    if meta_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        parts = []
        for k, v in meta_data.items():
            parts.append(f'{k}: {v}')
        run = p.add_run('  |  '.join(parts))
        run.font.size = Pt(9); run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_TEXT_MUTED

    # ── Fase badge ──
    if fase:
        p_fase = doc.add_paragraph()
        p_fase.paragraph_format.space_before = Pt(0)
        p_fase.paragraph_format.space_after  = Pt(10)
        run = p_fase.add_run(f'◆  {fase}')
        run.font.size = Pt(8); run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_ACCENT

    # ── Separador decorativo ──
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after  = Pt(8)
    _set_paragraph_border(p_sep, bottom_color="D4A030")
    run = p_sep.add_run('')
    run.font.size = Pt(2)

    # ── Procesar CARDS ──
    cards = soup.select('.card')
    for card in cards:
        _procesar_card(doc, card)

    # ── Salto de página ──
    doc.add_page_break()


def _procesar_card(doc, card):
    """Procesa un .card del HTML."""
    h2 = card.find('h2')
    titulo_card = h2.get_text(strip=True) if h2 else ''

    # FODA grid
    foda_grid = card.select_one('.foda-grid')
    if foda_grid:
        if titulo_card:
            doc.add_heading(titulo_card, level=2)
        _procesar_foda(doc, foda_grid)
        return

    # Tabla HTML
    table_container = card.select_one('.table-container')
    if table_container:
        if titulo_card:
            doc.add_heading(titulo_card, level=2)
        _procesar_tabla(doc, table_container.find('table'))
        return

    # Timeline
    timeline = card.select_one('.timeline')
    if timeline:
        if titulo_card:
            doc.add_heading(titulo_card, level=2)
        _procesar_timeline(doc, timeline)
        return

    # Key-value grid
    kv_grid = card.select_one('.kv-grid')
    if kv_grid:
        if titulo_card:
            doc.add_heading(titulo_card, level=2)
        _procesar_kv_grid(doc, kv_grid)
        return

    # Cards anidadas (arquetipos)
    sub_cards = card.select('.card')
    if sub_cards:
        # El título del padre ya es un heading, no repetir
        for sub in sub_cards:
            sub_h2 = sub.find('h2')
            if sub_h2:
                doc.add_heading(sub_h2.get_text(strip=True), level=3)
            sub_kv = sub.select_one('.kv-grid')
            if sub_kv:
                _procesar_kv_grid(doc, sub_kv)
            sub_foda = sub.select_one('.foda-grid')
            if sub_foda:
                _procesar_foda(doc, sub_foda)
        return

    # Fallback texto
    if titulo_card:
        doc.add_heading(titulo_card, level=2)

    p_text = card.find('p')
    if p_text:
        _make_paragraph(doc, p_text.get_text(strip=True), size=11,
                        color=COLOR_TEXT, space_after=10)

    ul = card.find('ul')
    if ul:
        for li in ul.find_all('li'):
            p = doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            for r in p.runs:
                r.font.size = Pt(10); r.font.name = FONT_BODY
                r.font.color.rgb = COLOR_TEXT


# ═══════════════════════════════════════════════════════════════
# KV GRID — Tabla elegante de 2 columnas
# ═══════════════════════════════════════════════════════════════

def _procesar_kv_grid(doc, kv_grid):
    """Convierte .kv-grid en tabla Word estilizada."""
    items = kv_grid.select('.kv-item')
    if not items:
        return

    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _aplicar_bordes_tabla(table, size='4', color='D0D5DD')

    for i, item in enumerate(items):
        label = item.select_one('.label')
        value = item.select_one('.value')
        lt = label.get_text(strip=True) if label else ''
        vt = value.get_text(strip=True) if value else ''

        cl = table.cell(i, 0)
        cv = table.cell(i, 1)

        _make_cell(cl, lt, bold=True, color=COLOR_PRIMARY_DARK, size=10)
        _make_cell(cv, vt, size=10, color=COLOR_TEXT)

        cl.width = Cm(5)
        cv.width = Cm(11)

        if i % 2 == 0:
            _set_cell_shading(cl, "F0F3F8")
        else:
            _set_cell_shading(cl, "FFFFFF")
            _set_cell_shading(cv, "FAFBFC")

    doc.add_paragraph('')


# ═══════════════════════════════════════════════════════════════
# TABLA HTML → Tabla Word estilizada
# ═══════════════════════════════════════════════════════════════

def _procesar_tabla(doc, table_html):
    """Convierte tabla HTML en tabla Word profesional."""
    if not table_html:
        return
    rows_html = table_html.find_all('tr')
    if not rows_html:
        return

    headers = [th.get_text(strip=True) for th in rows_html[0].find_all('th')]
    data_rows = []
    for tr in rows_html[1:]:
        tds = tr.find_all('td')
        row = [td.get_text(strip=True) for td in tds]
        if row: data_rows.append(row)

    ncols = max(len(headers), max((len(r) for r in data_rows), default=0))
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _aplicar_bordes_tabla(table, size='4', color='C0C5CD')

    # Header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _make_cell(cell, h, bold=True, color=COLOR_WHITE, size=9,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_shading(cell, HEX_TABLE_HEADER)

    # Data
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            if j < ncols:
                cell = table.cell(i + 1, j)
                _make_cell(cell, val, size=9, color=COLOR_TEXT)
                if i % 2 == 1:
                    _set_cell_shading(cell, "F5F7FA")

    doc.add_paragraph('')


# ═══════════════════════════════════════════════════════════════
# FODA — Cuadrícula 2×2 con colores
# ═══════════════════════════════════════════════════════════════

def _procesar_foda(doc, foda_grid):
    """Convierte FODA en tabla 2×2 con cuadrantes coloreados."""
    foda_cards = foda_grid.select('.foda-card')
    if not foda_cards:
        return

    # Mapeo de tipos: (título, color_rgb, hex_fondo, hex_borde)
    config = {
        'foda-fortaleza':   ('◆  Fortalezas',  COLOR_F_FORTALEZA,   'E6F7EE', 'A7F3D0'),
        'foda-debilidad':   ('●  Debilidades', COLOR_F_DEBILIDAD,   'FDE8E8', 'FECACA'),
        'foda-oportunidad': ('▸  Oportunidades', COLOR_F_OPORTUNIDAD, 'EBF3FE', 'BFDBFE'),
        'foda-amenaza':     ('✦  Amenazas',    COLOR_F_AMENAZA,     'FFF3E0', 'FED7AA'),
    }

    # Buscar en orden F1, D1, O1, A1 — es decir, los 4 en el orden que aparecen
    items = []
    for card in foda_cards:
        cls = card.get('class', [])
        tipo = next((c for c in cls if c in config), None)
        if tipo:
            titulo, color, bg_hex, border_hex = config[tipo]
            ul = card.find('ul')
            bullets = []
            if ul:
                for li in ul.find_all('li'):
                    t = li.get_text(strip=True)
                    t = re.sub(r'^[✅]\s*', '', t)
                    bullets.append(t)
            items.append((titulo, color, bg_hex, border_hex, bullets))

    if not items:
        return

    # Crear tabla 2×2
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _aplicar_bordes_tabla(table, size='8', color='D0D5DD', clean_first=True)

    # Llenar celdas
    for idx, (titulo, color, bg_hex, border_hex, bullets) in enumerate(items):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)

        # Limpiar y preparar
        cell.text = ''
        _set_cell_shading(cell, bg_hex)
        _set_cell_vertical_alignment(cell, 'top')

        # Título del cuadrante
        p_tit = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p_tit.paragraph_format.space_before = Pt(6)
        p_tit.paragraph_format.space_after  = Pt(4)
        r = p_tit.add_run(titulo)
        r.bold = True; r.font.size = Pt(11); r.font.name = FONT_TITLE
        r.font.color.rgb = color

        # Ancho de columna
        cell.width = Cm(8)

        # Items
        for b in bullets:
            p_b = cell.add_paragraph()
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after  = Pt(1)
            p_b.paragraph_format.left_indent  = Cm(0.5)
            p_b.style = doc.styles['List Bullet']
            # Quitar el existing text y agregar run
            p_b.clear()
            r_b = p_b.add_run(b)
            r_b.font.size = Pt(9); r_b.font.name = FONT_BODY
            r_b.font.color.rgb = COLOR_TEXT

    doc.add_paragraph('')


# ═══════════════════════════════════════════════════════════════
# TIMELINE — Tabla de objetivos
# ═══════════════════════════════════════════════════════════════

def _procesar_timeline(doc, timeline):
    """Convierte timeline en tabla Word."""
    items = timeline.select('.timeline-item')
    if not items:
        return

    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _aplicar_bordes_tabla(table, size='4', color='D0D5DD')

    for i, item in enumerate(items):
        title_el = item.select_one('.tl-title')
        desc_el  = item.select_one('.tl-desc')
        title = title_el.get_text(strip=True) if title_el else ''
        desc  = desc_el.get_text(strip=True)  if desc_el  else ''

        ct = table.cell(i, 0)
        cd = table.cell(i, 1)

        _make_cell(ct, title, bold=True, color=COLOR_PRIMARY_DARK, size=10)
        _make_cell(cd, desc, size=10, color=COLOR_TEXT)

        ct.width = Cm(4)
        cd.width = Cm(12)

        if i % 2 == 0:
            _set_cell_shading(ct, "F0F3F8")
        else:
            _set_cell_shading(ct, "FFFFFF")
            _set_cell_shading(cd, "FAFBFC")

    doc.add_paragraph('')


# ═══════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ═══════════════════════════════════════════════════════════════

def generar_word(cliente: str, output_path: str | None = None) -> str | None:
    """
    Genera un documento Word con todos los reportes HTML del cliente.
    Retorna la ruta del archivo generado.
    """
    carpeta_cliente = os.path.join(RUTA_CLIENTES, cliente)
    if not os.path.isdir(carpeta_cliente):
        logger.error('❌ No se encontró la carpeta para "%s".', cliente)
        logger.error('   Ruta esperada: %s', carpeta_cliente)
        return None

    pattern = os.path.join(carpeta_cliente, '*.html')
    all_htmls = [f for f in glob(pattern) if not os.path.basename(f).startswith('_')]
    if not all_htmls:
        logger.error('❌ No se encontraron reportes HTML en %s', carpeta_cliente)
        return None

    def _sort_key(f: str) -> tuple[int, int]:
        base = os.path.basename(f)
        m = re.match(r'(\d+)-(\d+)', base)
        if m: return (int(m.group(1)), int(m.group(2)))
        return (99, 99)

    all_htmls.sort(key=_sort_key)

    doc = Document()
    _crear_estilos(doc)

    # Extraer datos del primer HTML
    analista = 'Analista'
    fecha = datetime.now().strftime('%d/%m/%Y')
    if all_htmls:
        try:
            with open(all_htmls[0], 'r', encoding='utf-8') as f:
                s0 = BeautifulSoup(f.read(), 'html.parser')
            date_el = s0.select_one('.cover-date')
            if date_el:
                txt = date_el.get_text(strip=True)
                parts = txt.split('·')
                if len(parts) >= 2:
                    fecha = parts[0].strip()
                    analista = parts[1].strip()
        except Exception as e:
            logger.warning('  ⚠️ No se pudieron extraer metadatos de %s: %s', all_htmls[0], e)

    # ── Portada ──
    _add_cover(doc, cliente, analista, fecha, len(all_htmls))

    # ── Header y Footer para todas las secciones ──
    _add_header_footer(doc, cliente)

    # ── Procesar HTMLs ──
    for html_path in all_htmls:
        nombre = os.path.basename(html_path)
        logger.info('  Procesando: %s', nombre)

        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

        cover_h1 = soup.select_one('.cover-page h1')
        titulo = cover_h1.get_text(strip=True) if cover_h1 else 'Sin título'

        header_badge = soup.select_one('.header-badge')
        fase = header_badge.get_text(strip=True) if header_badge else ''

        paso = ''
        m = re.match(r'(\d+-\d+)', nombre)
        if m:
            paso = m.group(1)
        cover_sub = soup.select_one('.cover-page .cover-sub')
        if cover_sub and 'Paso' in cover_sub.get_text():
            paso = cover_sub.get_text(strip=True).replace('Paso', '').strip()

        _insert_seccion(doc, soup, paso, titulo, fase, cliente)

    # ── Guardar ──
    if output_path is None:
        fecha_archivo = datetime.now().strftime('%Y%m%d')
        output_path = os.path.join(carpeta_cliente, f'Onboarding-{cliente}-{fecha_archivo}.docx')

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def main() -> int:
    parser = argparse.ArgumentParser(
        description='Genera documento Word profesional con reportes de onboarding (diseño premium).'
    )
    parser.add_argument('--cliente', '-c', required=True,
                        help='Nombre del cliente (carpeta en Clientes/)')
    parser.add_argument('--output', '-o',
                        help='Ruta de salida del .docx (opcional)')
    args = parser.parse_args()

    logger.info('📄 Generando documento Word premium para: %s', args.cliente)
    logger.info('   Buscando reportes en: %s', os.path.join(RUTA_CLIENTES, args.cliente))
    logger.info('')

    ruta = generar_word(args.cliente, args.output)

    if ruta:
        logger.info('')
        logger.info('✅ Documento Word generado exitosamente:')
        logger.info('   📄 %s', ruta)
        return 0
    else:
        logger.info('')
        logger.error('❌ No se pudo generar el documento.')
        return 1


if __name__ == '__main__':
    sys.exit(main())
