#!/usr/bin/env python3
"""
Generador de documento Word - Metodología Integrada Global Infinity Marketing
Fusión: Metodología del usuario + Enfoque operativo con MCPs

Uso:
    python generar_metodologia.py
    python generar_metodologia.py --output "ruta/personalizada.docx"
"""

import argparse
import datetime
import logging
import os
import sys
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


def _build_document() -> Document:
    doc = Document()

    # Colores corporativos
    DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
    MEDIUM_BLUE = RGBColor(0x2D, 0x5F, 0x8A)
    ACCENT_BLUE = RGBColor(0x3A, 0x7B, 0xD5)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
    GOLD = RGBColor(0xD4, 0xA0, 0x1E)

    # ── Configuración de estilos ─────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Funciones auxiliares ──────────────────────────────────────────

    def add_heading_styled(text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                run.font.color.rgb = DARK_BLUE
            elif level == 2:
                run.font.color.rgb = MEDIUM_BLUE
            elif level == 3:
                run.font.color.rgb = ACCENT_BLUE
        return heading

    def add_paragraph_styled(text: str, bold: bool = False, italic: bool = False, color: RGBColor = DARK_GRAY, size: int = 11, alignment: Optional[int] = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        if alignment:
            p.alignment = alignment
        return p

    def add_bullet(text: str, level: int = 0, bold_prefix: str = ""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.size = Pt(11)
            run_b.font.color.rgb = DARK_GRAY
            run_n = p.add_run(text)
            run_n.font.size = Pt(11)
            run_n.font.color.rgb = DARK_GRAY
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
        return p

    def add_table_with_data(headers: list, rows: list, col_widths: Optional[list] = None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = WHITE
            # Dark blue background
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B2A4A"/>')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row_data):
                row_cells[c_idx].text = str(cell_text)
                for paragraph in row_cells[c_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = DARK_GRAY
            # Alternate row shading
            if r_idx % 2 == 1:
                for c_idx in range(len(headers)):
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FC"/>')
                    row_cells[c_idx]._tc.get_or_add_tcPr().append(shading_elm)

        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)

        doc.add_paragraph()  # spacer
        return table

    def add_step_box(step_num: str, step_title: str, description: str, tools: str, deliverable: str):
        heading = doc.add_heading(f"Paso {step_num}: {step_title}", level=3)
        for run in heading.runs:
            run.font.color.rgb = ACCENT_BLUE

        p_desc = doc.add_paragraph()
        run_label = p_desc.add_run("📋 Descripción: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_text = p_desc.add_run(description)
        run_text.font.size = Pt(11)

        p_tools = doc.add_paragraph()
        run_label = p_tools.add_run("🛠️ Herramientas: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_text = p_tools.add_run(tools)
        run_text.font.size = Pt(11)

        p_deliverable = doc.add_paragraph()
        run_label = p_deliverable.add_run("📦 Entregable: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_text = p_deliverable.add_run(deliverable)
        run_text.font.size = Pt(11)

        doc.add_paragraph()  # spacer

    def add_checklist(items: list):
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run(f"☐  {item}")
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
        doc.add_paragraph()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║                       PORTADA                                   ║
    # ╚══════════════════════════════════════════════════════════════════╝

    # Sección de portada en horizontal
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Espaciado superior
    for _ in range(6):
        doc.add_paragraph()

    # Logo / Marca
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("GLOBAL INFINITY MARKETING")
    run_logo.bold = True
    run_logo.font.size = Pt(28)
    run_logo.font.color.rgb = DARK_BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Metodología Integrada de Trabajo")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = MEDIUM_BLUE

    # Línea decorativa
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("─" * 50)
    run_line.font.color.rgb = GOLD

    # Subtítulo
    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub2 = p_sub2.add_run("Proceso completo de onboarding, diagnóstico y ejecución\npara campañas de marketing digital en Uruguay")
    run_sub2.font.size = Pt(13)
    run_sub2.font.color.rgb = DARK_GRAY
    run_sub2.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Info
    today = datetime.date.today().strftime("%d de %B de %Y")
    info_lines = [
        f"Versión: 1.0 — {today}",
        "Documento interno — Global Infinity Marketing",
        "Confidencial"
    ]
    for line in info_lines:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_info = p_info.add_run(line)
        run_info.font.size = Pt(11)
        run_info.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║                      ÍNDICE                                     ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("Índice", level=1)

    toc_items = [
        ("1.", "Metodología Integrada — Visión General"),
        ("2.", "Fase 1 — Diagnóstico Estratégico"),
        ("  2.1", "Entrevista de Descubrimiento"),
        ("  2.2", "Análisis del Producto / Servicio"),
        ("  2.3", "Definición del Cliente Ideal (ICP)"),
        ("  2.4", "Identidad de Marca"),
        ("  2.5", "Definición de Objetivos SMART"),
        ("  2.6", "Análisis Interno (FODA)"),
        ("3.", "Fase 2 — Investigación de Mercado"),
        ("  3.1", "Análisis de Competidores"),
        ("  3.2", "Tendencias y Estacionalidad"),
        ("  3.3", "Benchmark de Precios y Promociones"),
        ("  3.4", "Investigación de Palabras Clave"),
        ("4.", "Fase 3 — Auditoría de Activos Digitales"),
        ("  4.1", "Auditoría de Sitio Web / E-commerce"),
        ("  4.2", "Auditoría de Redes Sociales"),
        ("  4.3", "Revisión de Tracking y Píxeles"),
        ("  4.4", "Análisis SEO (Google Search Console)"),
        ("  4.5", "Diagnóstico de Email Marketing"),
        ("5.", "Fase 4 — Presencia Digital"),
        ("  5.1", "Desarrollo o Mejora del Sitio Web"),
        ("  5.2", "Optimización de Perfiles Sociales"),
        ("  5.3", "Configuración de Herramientas de Medición"),
        ("6.", "Fase 5 — Estrategia y Plan de Medios"),
        ("  6.1", "Desarrollo de la Propuesta Comercial"),
        ("  6.2", "Definición de Canales y Presupuesto"),
        ("  6.3", "KPIs por Canal y Proyecciones"),
        ("7.", "Fase 6 — Setup Técnico y Creatividades"),
        ("  7.1", "Configuración de Píxeles y Tags"),
        ("  7.2", "Conversiones en Google Ads"),
        ("  7.3", "Feed de Productos (Google Shopping)"),
        ("  7.4", "Creación de Audiencias y Segmentos"),
        ("  7.5", "Diseño de Creatividades"),
        ("8.", "Fase 7 — Pilotaje y Testeo"),
        ("  8.1", "Campaña Piloto (7-14 días)"),
        ("  8.2", "Pruebas A/B"),
        ("  8.3", "Análisis de Resultados y Ajustes"),
        ("9.", "Fase 8 — Lanzamiento y Escalado"),
        ("  9.1", "Escalar Campañas Exitosas"),
        ("  9.2", "Marketing de Contenidos (Paralelo)"),
        ("  9.3", "Automatización de Reportes"),
        ("  9.4", "Optimización Continua"),
        ("10.", "Anexos — Tabla de Herramientas MCP"),
        ("11.", "Checklist de Proyecto"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        if num.startswith(" "):
            p.paragraph_format.left_indent = Cm(1.5)
        else:
            p.paragraph_format.space_before = Pt(4)
        run_num = p.add_run(f"{num}  ")
        run_num.bold = True
        run_num.font.size = Pt(11)
        run_num.font.color.rgb = DARK_BLUE
        run_title = p.add_run(title)
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║            1. METODOLOGÍA INTEGRADA — VISIÓN GENERAL            ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("1. Metodología Integrada — Visión General", level=1)

    add_paragraph_styled(
        "Esta metodología integra el enfoque estratégico-conceptual con el plan operativo-detallado "
        "para garantizar que cada campaña publicitaria se base en un entendimiento profundo del negocio, "
        "el mercado uruguayo y los datos reales del cliente."
    )

    add_heading_styled("Estructura General", level=2)

    add_table_with_data(
        ["Fase", "Nombre", "Duración Estimada", "Objetivo Principal"],
        [
            ["1", "Diagnóstico Estratégico", "3-5 días", "Entender el negocio, su identidad y objetivos"],
            ["2", "Investigación de Mercado", "2-3 días", "Analizar competencia, tendencias y keywords"],
            ["3", "Auditoría de Activos Digitales", "2-4 días", "Evaluar sitio web, tracking, SEO y redes"],
            ["4", "Presencia Digital", "5-15 días", "Optimizar o crear activos digitales base"],
            ["5", "Estrategia y Plan de Medios", "2-3 días", "Definir canales, presupuesto y KPIs"],
            ["6", "Setup Técnico y Creatividades", "3-5 días", "Configurar píxeles, audiencias y anuncios"],
            ["7", "Pilotaje y Testeo", "7-14 días", "Validar campaña con inversión controlada"],
            ["8", "Lanzamiento y Escalado", "Continuo", "Escalar, automatizar y optimizar"],
        ],
        col_widths=[1.5, 5.5, 3.5, 7.5]
    )

    add_paragraph_styled(
        "Duración total estimada del onboarding completo: 3 a 6 semanas, dependiendo de la complejidad "
        "del cliente y el estado de sus activos digitales.",
        italic=True, size=10
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 1 — DIAGNÓSTICO ESTRATÉGICO                        ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("2. Fase 1 — Diagnóstico Estratégico", level=1)

    add_paragraph_styled(
        "Objetivo: Conocer el negocio en profundidad para garantizar que todas las acciones de marketing "
        "atraigan oportunidades comerciales relevantes y alineadas con la realidad de la empresa."
    )

    # 2.1
    add_step_box(
        "2.1", "Entrevista de Descubrimiento",
        "Reunión inicial con el cliente para relevar información fundamental del negocio. "
        "Se aborda: historia de la empresa, productos/servicios, modelo de negocio, canales de venta actuales, "
        "proceso de atención al cliente, público actual vs deseado, diferenciación competitiva, y objetivos de negocio.",
        "Google Meet / Zoom, Notion (toma de notas), Google Docs (brief colaborativo)",
        "Brief de negocio completo documentado"
    )

    add_heading_styled("Checklist de preguntas clave", level=3)
    add_checklist([
        "¿Cuál es la historia de la empresa? ¿Cuánto tiempo llevan en el mercado?",
        "¿Qué productos o servicios venden? ¿Cuáles son los más rentables?",
        "¿Cuál es su modelo de negocio? (B2B, B2C, suscripción, venta directa, etc.)",
        "¿Quiénes son sus clientes actuales? ¿Quiénes les gustaría que fueran?",
        "¿Qué los diferencia de la competencia?",
        "¿Cuáles son sus canales de venta actuales? (físico, web, Mercado Libre, redes, etc.)",
        "¿Cómo es el proceso de atención al cliente? ¿Hay posventa?",
        "¿Cuáles son sus objetivos a 3, 6 y 12 meses?",
        "¿Han hecho publicidad digital antes? ¿Qué funcionó y qué no?",
        "¿Cuál es el presupuesto mensual estimado para publicidad?"
    ])

    # 2.2
    add_step_box(
        "2.2", "Análisis del Producto / Servicio",
        "Catalogar y analizar cada producto o servicio: precios, márgenes, categorías, estacionalidad, "
        "stock disponible, diferenciales frente a competencia. Identificar productos 'estrella' (alto margen + alta demanda) "
        "y productos 'imán' (bajo precio para atraer tráfico).",
        "Mercado Libre (comparar precios UY), Google Sheets (catalogación), Google Search (investigación)",
        "Matriz de productos con precios, márgenes y diferenciales"
    )

    # 2.3
    add_step_box(
        "2.3", "Definición del Cliente Ideal (ICP)",
        "Construir el perfil del cliente ideal basado en datos reales (si existen) o hipótesis validadas. "
        "Incluir: edad, ubicación geográfica (Uruguay por departamentos), nivel de ingresos, intereses, "
        "comportamiento de compra online, dispositivos usados, horarios de conexión, objeciones de compra.",
        "Brave Search / Google Search (datos demográficos UY), Meta Ads (audience insights si ya hay datos), GA4 (audiencias existentes)",
        "Perfil de Cliente Ideal (ICP) documentado con 1-3 arquetipos"
    )

    # 2.4
    add_step_box(
        "2.4", "Identidad de Marca",
        "Definir o validar los pilares de la marca que condicionarán el tono y contenido de toda la comunicación: "
        "Misión, Visión, Valores, Tono de comunicación, Posicionamiento deseado, e Imagen corporativa actual. "
        "Si el cliente ya los tiene definidos, se validan. Si no, se construyen.",
        "Google Docs (documento colaborativo), Canva (referencias visuales), Notion (brief de marca)",
        "Documento de Identidad de Marca (Misión, Visión, Valores, Tono, Posicionamiento)"
    )

    # 2.5
    add_step_box(
        "2.5", "Definición de Objetivos SMART",
        "Convertir los objetivos de negocio del cliente en metas específicas, medibles, alcanzables, relevantes y con plazo. "
        "Ejemplo: 'Aumentar ventas online' → 'Generar 50 ventas adicionales por mes con un ROAS ≥ 4x en los próximos 60 días'. "
        "Tipos de objetivos: conseguir más clientes, vender más a clientes actuales, lanzar un producto, expandirse, reconocimiento de marca.",
        "Sequential Thinking (análisis estructurado), Google Sheets (seguimiento de metas)",
        "Documento de Objetivos SMART con KPIs primarios y secundarios"
    )

    # 2.6
    add_step_box(
        "2.6", "Análisis Interno (FODA)",
        "Evaluar capacidades internas del cliente: Fortalezas (recursos, equipo, tecnología), Debilidades (limitaciones, falta de personal, tecnología obsoleta), "
        "Oportunidades (mercado, tendencias, nichos desatendidos en Uruguay), Amenazas (competencia, cambios regulatorios, situación económica). "
        "Evaluar también capacidad de atención al cliente (¿pueden manejar un aumento de demanda?) y presupuesto disponible.",
        "Google Docs (matriz FODA), Sequential Thinking (análisis estratégico)",
        "Matriz FODA del negocio con conclusiones accionables"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 2 — INVESTIGACIÓN DE MERCADO                       ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("3. Fase 2 — Investigación de Mercado", level=1)

    add_paragraph_styled(
        "Objetivo: Analizar el entorno competitivo y las oportunidades del mercado uruguayo para "
        "fundamentar la estrategia con datos reales."
    )

    # 3.1
    add_step_box(
        "3.1", "Análisis de Competidores",
        "Identificar 3 a 5 competidores directos en Uruguay. Analizar: qué canales usan (Google Ads, Meta, TikTok, ML), "
        "qué tipo de anuncios tienen (Meta Ads Library), qué palabras clave pujan (Google Search), "
        "qué precios y promociones manejan, cuál es su propuesta de valor, reputación online y reseñas en Mercado Libre. "
        "Además, revisar su presencia en redes y engagement.",
        "Brave Search, Google Search, Mercado Libre (productos, reviews, reputación), Meta Ads Library, Google Search Console (si tenemos acceso)",
        "Matriz competitiva con fortalezas, debilidades y estrategias de cada competidor"
    )

    # 3.2
    add_step_box(
        "3.2", "Tendencias y Estacionalidad",
        "Identificar picos de demanda, eventos comerciales y estacionalidad del mercado uruguayo: "
        "Hot Sale (julio/agosto), Cyber Monday (noviembre), Día del Padre (junio), Día de la Madre (mayo), "
        "Navidad y Reyes (diciembre-enero), Semana de Turismo (marzo/abril), Día del Niño (agosto), "
        "vuelta a clases (febrero/marzo), temporada turística (enero). También tendencias sectoriales específicas.",
        "Brave Search, Google Search (tendencias), Google Trends, Sequential Thinking",
        "Calendario de oportunidades comerciales con fechas clave y recomendaciones"
    )

    # 3.3
    add_step_box(
        "3.3", "Benchmark de Precios y Promociones",
        "Relevamiento detallado de estrategias de precios de competidores: rangos de precio por producto, "
        "promociones recurrentes (2x1, descuentos por volumen, envío gratis), "
        "estrategias de financiación (Cuotas sin interés en Uruguay), y precio psicológico. "
        "Especial foco en Mercado Libre Uruguay como termómetro de precios.",
        "Mercado Libre (comparativa de precios y promociones), Google Search (ofertas visibles)",
        "Reporte de benchmark de precios con insights accionables"
    )

    # 3.4
    add_step_box(
        "3.4", "Investigación de Palabras Clave",
        "Identificar las palabras clave que usa el público uruguayo para buscar los productos/servicios del cliente. "
        "Analizar: volumen de búsqueda, intención (informativa, comercial, transaccional), estacionalidad, "
        "y nivel de competencia. Priorizar keywords long-tail con alta intención de compra. "
        "Incluir variantes locales uruguayas (ej: 'precio', 'comprar', 'envío gratis', 'en Uruguay').",
        "Google Search Console (datos reales del cliente si tiene), Google Ads Keyword Planner, Brave Search, Google Search",
        "Lista de palabras clave priorizadas por intención y oportunidad"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 3 — AUDITORÍA DE ACTIVOS DIGITALES                 ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("4. Fase 3 — Auditoría de Activos Digitales", level=1)

    add_paragraph_styled(
        "Objetivo: Evaluar el estado actual de todos los activos digitales del cliente para identificar "
        "brechas, problemas técnicos y oportunidades de mejora antes de invertir en publicidad."
    )

    # 4.1
    add_step_box(
        "4.1", "Auditoría de Sitio Web / E-commerce",
        "Revisión exhaustiva del sitio web: velocidad de carga (Lighthouse), optimización mobile, "
        "experiencia de usuario (UX), claridad del funnel de conversión, proceso de carrito y checkout, "
        "políticas de envío y devolución visibles, formularios de contacto/lead, "
        "seguridad (SSL), y accesibilidad. Si es e-commerce, revisar pasarela de pago y opciones de envío Uruguay.",
        "Chrome DevTools (Lighthouse, rendimiento), GA4 (métricas de comportamiento), Google Search Console (indexación)",
        "Informe de salud del sitio con problemas detectados y prioridad de solución"
    )

    # 4.2
    add_step_box(
        "4.2", "Auditoría de Redes Sociales",
        "Evaluar la presencia actual del cliente en redes sociales: seguidores, tasa de engagement, "
        "frecuencia de publicación, calidad del contenido, tono de comunicación, "
        "respuesta a comentarios/mensajes, tipo de contenido que mejor funciona. "
        "Incluir Instagram, Facebook, X/Twitter, LinkedIn, TikTok y YouTube según aplique.",
        "X/Twitter (análisis de perfil), Meta Ads (biblioteca de anuncios, insights), Brave Search (tendencias redes UY)",
        "Diagnóstico de redes sociales con métricas actuales y recomendaciones"
    )

    # 4.3
    add_step_box(
        "4.3", "Revisión de Tracking y Píxeles",
        "PUNTO CRÍTICO. Verificar la correcta instalación de: Meta Pixel (eventos estándar y personalizados), "
        "Google tag (Google Ads + GA4), Google Analytics 4 (propiedad, streams, eventos), "
        "Google Search Console (sitio verificado). Validar que los eventos de conversión estén disparando correctamente: "
        "compra, add_to_cart, initiate_checkout, lead, formulario, etc. Revisar atribución y deduplicación de conversiones.",
        "GA4 (debug view, eventos en tiempo real), Meta Ads (píxel, eventos), Google Ads (conversiones), Google Search Console",
        "Diagnóstico técnico de tracking con lista de correcciones necesarias"
    )

    # 4.4
    add_step_box(
        "4.4", "Análisis SEO (Google Search Console)",
        "Analizar rendimiento orgánico actual: clics, impresiones, CTR, posición media en Google. "
        "Identificar: páginas con más tráfico, consultas que generan impresiones pero bajo CTR (oportunidades de título/meta), "
        "posiciones 4-15 con potencial de subir a página 1 (quick wins), "
        "pérdidas de tráfico recientes, canibalización de keywords, contenido decayente. "
        "Revisar sitemaps y estado de indexación.",
        "Google Search Console (site snapshot, quick wins, content gaps, traffic drops, alerts, cannibalization)",
        "Reporte SEO completo con quick wins priorizados"
    )

    # 4.5
    add_step_box(
        "4.5", "Diagnóstico de Email Marketing",
        "Si el cliente ya tiene lista de contactos o usa email marketing: revisar tamaño y calidad de la lista, "
        "segmentación actual, tasas de apertura y clics, flows automáticos activos (bienvenida, abandono, postventa), "
        "diseño de emails, y herramientas utilizadas.",
        "Klaviyo (flows, campañas, segmentos, métricas)",
        "Diagnóstico de email marketing con recomendaciones de mejora"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 4 — PRESENCIA DIGITAL                              ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("5. Fase 4 — Presencia Digital", level=1)

    add_paragraph_styled(
        "Objetivo: Asegurar que el cliente tenga una presencia digital sólida y profesional antes de "
        "invertir en tráfico pago. No tiene sentido llevar tráfico a un sitio lento o perfiles descuidados."
    )

    # 5.1
    add_step_box(
        "5.1", "Desarrollo o Mejora del Sitio Web",
        "Según los hallazgos de la auditoría: implementar mejoras de velocidad, optimización mobile, "
        "claridad del mensaje principal, llamados a la acción (CTAs) visibles, formularios funcionales, "
        "políticas de envío y devolución claras. Si no tiene sitio, desarrollar uno mínimo viable (landing page + catálogo). "
        "Si usa Shopify / Tiendanube / WooCommerce, optimizar plantilla y configuración.",
        "Shopify Expert, WordPress Pro, Google Search Console (validar indexación), GA4 (medir impacto)",
        "Sitio web optimizado y listo para recibir tráfico pago"
    )

    # 5.2
    add_step_box(
        "5.2", "Optimización de Perfiles Sociales",
        "Actualizar y optimizar perfiles en las redes sociales relevantes: bio con propuesta de valor clara, "
        "link en bio funcional (linktree o similar), imágenes de perfil y portada profesionales, "
        "destacados organizados (Instagram), información de contacto completa. "
        "Crear o completar perfiles que falten según la estrategia.",
        "Meta Ads (gestión de páginas), X/Twitter, Canva (diseño de perfiles)",
        "Perfiles sociales optimizados con identidad visual coherente"
    )

    # 5.3
    add_step_box(
        "5.3", "Configuración de Herramientas de Medición",
        "Si no están configuradas: instalar y configurar GA4 (propiedad, streams web/app, eventos), "
        "Google Search Console (verificar sitio, enviar sitemap), Meta Pixel (eventos estándar + personalizados), "
        "Google tag, y Klaviyo (si aplica email). Configurar dashboards iniciales en GA4 y Google Sheets.",
        "GA4, Google Search Console, Meta Ads, Google Ads, Klaviyo, Google Sheets (dashboard), n8n (automatización)",
        "Stack analítico completo y funcional"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 5 — ESTRATEGIA Y PLAN DE MEDIOS                   ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("6. Fase 5 — Estrategia y Plan de Medios", level=1)

    add_paragraph_styled(
        "Objetivo: Definir la hoja de ruta de marketing digital con canales, presupuesto, mensajes y métricas, "
        "basándose en toda la información recopilada en las fases anteriores."
    )

    # 6.1
    add_step_box(
        "6.1", "Desarrollo de la Propuesta Comercial",
        "Formalizar la propuesta de valor única (UVP) que se comunicará en todos los canales. "
        "Definir: mensaje principal, argumentos de venta, objeciones a abordar, "
        "y diferenciación clara frente a competidores. Crear el elevator pitch del negocio.",
        "Sequential Thinking (razonamiento estratégico), Google Docs (documentación)",
        "Propuesta comercial documentada con UVP y mensajes clave"
    )

    # 6.2
    add_step_box(
        "6.2", "Definición de Canales y Presupuesto",
        "Seleccionar los canales publicitarios según el tipo de negocio y objetivos: "
        "Meta Ads (Facebook + Instagram) para branding, interés y remarketing. "
        "Google Ads (Search + Shopping) para captar intención de compra activa. "
        "Distribución recomendada: \n"
        "• Empresas nuevas / en descubrimiento: 70% Meta Ads, 30% Google Ads.\n"
        "• Empresas con demanda existente: 50% Meta Ads, 50% Google Ads.\n"
        "Asignar presupuesto mensual por canal y establecer frecuencia de revisión.",
        "Google Ads (planificador de presupuesto), Meta Ads (estimaciones de alcance), Google Sheets (planilla)",
        "Plan de medios con asignación presupuestal y justificación"
    )

    # 6.3
    add_step_box(
        "6.3", "KPIs por Canal y Proyecciones",
        "Definir métricas clave por canal y tipo de campaña:\n"
        "• Meta Ads: CPC, CPM, CTR, CPA, ROAS, frecuencia, tasa de engagement.\n"
        "• Google Ads: CTR, CPC, CPA, ROAS, impresiones, tasa de conversión, Quality Score.\n"
        "• SEO: clics, impresiones, CTR, posición media, leads orgánicos.\n"
        "• Email: tasa de apertura, CTR, tasa de conversión, ingresos por email.\n"
        "Establecer metas numéricas realistas y proyectar resultados esperados.",
        "Sequential Thinking (proyecciones), Google Sheets (modelo financiero), GA4 (benchmarks históricos)",
        "Tablero de KPIs con metas por canal y proyección de resultados"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 6 — SETUP TÉCNICO Y CREATIVIDADES                  ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("7. Fase 6 — Setup Técnico y Creatividades", level=1)

    add_paragraph_styled(
        "Objetivo: Preparar toda la infraestructura técnica y creativa para el lanzamiento de campañas."
    )

    # 7.1
    add_step_box(
        "7.1", "Configuración de Píxeles y Tags",
        "Instalar o verificar: Meta Pixel con eventos estándar (ViewContent, AddToCart, Purchase, Lead) "
        "y eventos personalizados según funnel del cliente. Google tag con conversiones importadas desde GA4. "
        "Configurar Google Tag Manager si es necesario para gestión centralizada. "
        "Realizar pruebas de disparo de eventos con Meta Pixel Helper y GA4 Debug View.",
        "Meta Ads (píxel, eventos), Google Ads (tag, conversiones), GA4 (debug view), Google Tag Manager",
        "Tracking funcional con eventos verificados"
    )

    # 7.2
    add_step_box(
        "7.2", "Conversiones en Google Ads",
        "Importar conversiones desde GA4 a Google Ads. Configurar value tracking (ingresos por conversión). "
        "Establecer ventana de atribución (recomendado: 7 días para clicks, 1 día para impresiones en Uruguay). "
        "Configurar modelos de atribución y deduplicar conversiones entre Meta y Google.",
        "Google Ads (conversiones, atribución), GA4 (eventos de conversión)",
        "Conversiones de Google Ads configuradas y validadas"
    )

    # 7.3
    add_step_box(
        "7.3", "Feed de Productos (Google Shopping)",
        "Si el cliente vende productos físicos: crear o validar el feed de productos para Google Merchant Center. "
        "Incluir: título, descripción, precio, disponibilidad, link de producto, imagen, categoría, GTIN/MPN. "
        "Configurar y optimizar campañas de Shopping. Revisar políticas de precios y envío.",
        "Google Ads (Merchant Center, Shopping campaigns), Google Sheets (feed), Mercado Libre (validar datos)",
        "Feed de productos activo en Google Merchant Center"
    )

    # 7.4
    add_step_box(
        "7.4", "Creación de Audiencias y Segmentos",
        "Configurar audiencias antes del lanzamiento:\n"
        "• Remarketing: visitantes del sitio web, add-to-cart sin compra, visitantes de páginas clave.\n"
        "• Audiencias similares (lookalikes): a partir de clientes existentes o leads.\n"
        "• Segmentos personalizados: por palabras clave, intención de compra, datos demográficos.\n"
        "• Listas de clientes (CRM) para targeting y exclusión.",
        "Meta Ads (audiencias, lookalikes), Google Ads (remarketing, audiencias personalizadas), Klaviyo (segmentos), HubSpot (listas CRM)",
        "Audiencias configuradas y listas para segmentar campañas"
    )

    # 7.5
    add_step_box(
        "7.5", "Diseño de Creatividades",
        "Desarrollar las piezas publicitarias según la identidad de marca y el tono definidos:\n"
        "• Meta Ads: imágenes, videos, carruseles, stories, copy para cada etapa del funnel (interés, consideración, conversión).\n"
        "• Google Ads: anuncios de texto (Responsive Search Ads), anuncios de Shopping, anuncios display.\n"
        "Crear mínimo 3-4 variantes por campaña para testear. Seguir lineamientos de Meta (20% texto en imagen) y Google.",
        "Canva / Photoshop (diseño), Meta Ads (biblioteca de anuncios para inspiración), Google Ads (RSA)",
        "Set completo de creatividades listas para pilotaje"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 7 — PILOTAJE Y TESTEO                              ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("8. Fase 7 — Pilotaje y Testeo", level=1)

    add_paragraph_styled(
        "Objetivo: Validar hipótesis con inversión controlada antes de escalar. "
        "Esta fase es obligatoria — nunca escalar sin datos propios del cliente."
    )

    # 8.1
    add_step_box(
        "8.1", "Campaña Piloto (7-14 días)",
        "Lanzar campañas con presupuesto reducido (recomendado: $200-$500 USD por canal) durante 7-14 días. "
        "Meta Ads: 2-3 conjuntos de anuncios testeando diferentes audiencias y creatividades. "
        "Google Ads: 2-3 campañas (Search + Shopping si aplica) con keywords priorizadas. "
        "Monitorear diariamente: CTR, CPC, CPA, tasa de conversión, frecuencia.",
        "Meta Ads (campañas en vivo), Google Ads (campañas en vivo), GA4 (tráfico y conversiones)",
        "Resultados de campaña piloto con datos reales"
    )

    # 8.2
    add_step_box(
        "8.2", "Pruebas A/B",
        "Durante el piloto, realizar pruebas A/B controladas:\n"
        "• Creatividades: imagen A vs imagen B, video vs imagen, copy A vs copy B.\n"
        "• Audiencias: interés A vs audiencia similar, remarketing vs prospección.\n"
        "• Ofertas: descuento vs envío gratis vs bonificación.\n"
        "• Landing pages: página A vs página B (si aplica).\n"
        "Una variable a la vez con muestra suficiente para significancia estadística.",
        "Meta Ads (pruebas A/B nativas), Google Ads (experimentos), Google Optimize (landing pages)",
        "Resultados de pruebas A/B con ganador identificado"
    )

    # 8.3
    add_step_box(
        "8.3", "Análisis de Resultados y Ajustes",
        "Evaluar qué funcionó y qué no según los KPIs definidos. Identificar:\n"
        "• Campañas/anuncios con mejor ROAS y CPA.\n"
        "• Audiencias más rentables.\n"
        "• Creatividades ganadoras.\n"
        "• Horarios y días con mejor rendimiento.\n"
        "• Dispositivos predominantes.\n"
        "Ajustar presupuestos: aumentar inversión en lo que funciona, pausar lo que no, "
        "y generar nuevas variantes basadas en los aprendizajes.",
        "GA4 (reportes), Meta Ads (insights), Google Ads (reportes), Sequential Thinking (análisis estratégico)",
        "Plan de optimización con acciones concretas"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        FASE 8 — LANZAMIENTO Y ESCALADO                         ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("9. Fase 8 — Lanzamiento y Escalado", level=1)

    add_paragraph_styled(
        "Objetivo: Escalar las campañas validadas, mantener el momentum con marketing de contenidos, "
        "automatizar reportes, y establecer un ciclo de optimización continua."
    )

    # 9.1
    add_step_box(
        "9.1", "Escalar Campañas Exitosas",
        "Aumentar presupuesto gradualmente (20-30% cada 3-4 días) en las campañas que cumplen KPIs. "
        "Expandir audiencias similares (lookalikes) de 1% a 2-3%, probar nuevos intereses, "
        "agregar nuevas keywords y ampliar segmentación geográfica dentro de Uruguay. "
        "No escalar más de una variable a la vez para mantener control.",
        "Meta Ads (escalado de presupuesto), Google Ads (ajuste de pujas y presupuestos)",
        "Campañas escaladas con monitoreo continuo"
    )

    # 9.2
    add_step_box(
        "9.2", "Marketing de Contenidos (Paralelo)",
        "Paralelo a la publicidad paga, desarrollar presencia orgánica:\n"
        "• Gestión de redes sociales con calendario editorial.\n"
        "• Creación de contenido relevante (posts, reels, stories, infografías).\n"
        "• SEO on-page: artículos de blog optimizados para keywords identificadas.\n"
        "• Posicionamiento de marca a través de contenido de valor.\n"
        "El contenido orgánico alimenta y abarata las campañas pagas (remarketing, audiencias calientes).",
        "Meta Ads (programación de contenido), X/Twitter (publicación), Google Search Console (SEO), Klaviyo (newsletter)",
        "Calendario editorial + contenido semanal"
    )

    # 9.3
    add_step_box(
        "9.3", "Automatización de Reportes",
        "Configurar reportes automáticos para no depender de revisión manual constante:\n"
        "• Dashboard semanal automatizado en Google Sheets con datos de Meta, Google, GA4 y Klaviyo.\n"
        "• Alertas automáticas (vía n8n) cuando KPIs caen por debajo de umbrales: "
        "CPA > límite, ROAS < mínimo, CTR anormalmente bajo.\n"
        "• Reporte ejecutivo mensual con insights y recomendaciones.",
        "n8n (automatización de flujos), Google Sheets (dashboard en vivo), Notion (documentación de campañas)",
        "Sistema de reporting automatizado con alertas"
    )

    # 9.4
    add_step_box(
        "9.4", "Optimización Continua",
        "Ciclo semanal de revisión y optimización:\n"
        "• Lunes: revisar resultados de la semana anterior, identificar anomalías.\n"
        "• Martes: ajustar campañas basado en datos (pujas, presupuestos, segmentación).\n"
        "• Miércoles: refrescar creatividades (evitar fatiga publicitaria).\n"
        "• Jueves: probar nuevas variantes (A/B testing continuo).\n"
        "• Viernes: preparar resumen semanal y plan para la semana siguiente.\n"
        "Revisión mensual profunda: redefinir estrategia si es necesario.",
        "GA4, Meta Ads, Google Ads, Klaviyo, Sequential Thinking, n8n (alertas automáticas)",
        "Ciclo semanal de optimización documentado"
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        ANEXOS — TABLA DE HERRAMIENTAS MCP                      ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("10. Anexos — Tabla de Herramientas MCP", level=1)

    add_paragraph_styled(
        "Las siguientes herramientas (MCP) están disponibles en Global Infinity Marketing "
        "para ejecutar cada fase de la metodología:"
    )

    add_table_with_data(
        ["Herramienta MCP", "Tipo", "¿Para qué se usa?", "Estado"],
        [
            ["Mercado Libre", "API", "Consultar productos, precios, competidores, categorías y reputación en MLU", "Requiere credenciales"],
            ["Google Ads", "API", "Gestionar campañas, presupuestos, keywords, conversiones y métricas", "Requiere credenciales"],
            ["Meta Ads", "API", "Administrar campañas de Facebook e Instagram Ads", "✅ Listo (OAuth)"],
            ["GA4", "API", "Analizar tráfico web, conversiones, audiencias y eventos", "Requiere credenciales"],
            ["HubSpot", "API", "Gestionar CRM, pipeline de ventas, contactos y deals", "Requiere credenciales"],
            ["Google Sheets", "API", "Crear y actualizar dashboards y reportes", "Requiere credenciales"],
            ["Klaviyo", "API", "Gestionar email marketing, flows, segmentos y campañas", "✅ Listo (OAuth)"],
            ["Brave Search", "API", "Búsqueda web de tendencias, competidores y noticias", "Requiere API key"],
            ["Google Search", "Local", "Búsqueda de información de mercado y competidores", "✅ Listo"],
            ["Sequential Thinking", "Local", "Razonamiento estructurado para análisis estratégico", "✅ Listo"],
            ["Google Search Console", "API", "Analizar SEO, posiciones, clics, impresiones y CTR", "Requiere credenciales"],
            ["X/Twitter", "Local", "Publicar tweets, buscar, analizar perfiles y tendencias", "✅ Listo"],
            ["Notion", "API", "Documentación de campañas, briefs y estrategias", "Requiere token"],
            ["n8n", "Local", "Automatizar flujos de trabajo, alertas, reportes", "Requiere instalación"],
        ],
        col_widths=[4.0, 1.5, 7.5, 3.5]
    )

    doc.add_page_break()

    # ╔══════════════════════════════════════════════════════════════════╗
    # ║        CHECKLIST DE PROYECTO                                   ║
    # ╚══════════════════════════════════════════════════════════════════╝

    add_heading_styled("11. Checklist de Proyecto", level=1)

    add_paragraph_styled(
        "Usar este checklist para tracking del progreso con cada cliente nuevo:",
        bold=True, size=12
    )

    checklist_items = [
        # Fase 1
        ("FASE 1 — DIAGNÓSTICO ESTRATÉGICO", True),
        ("☐  Entrevista de descubrimiento realizada y brief documentado", False),
        ("☐  Análisis de producto/servicio completado", False),
        ("☐  Perfil de Cliente Ideal (ICP) definido", False),
        ("☐  Identidad de marca documentada (Misión, Visión, Valores, Tono)", False),
        ("☐  Objetivos SMART definidos y acordados con el cliente", False),
        ("☐  Matriz FODA completada", False),
        # Fase 2
        ("FASE 2 — INVESTIGACIÓN DE MERCADO", True),
        ("☐  Matriz competitiva con 3-5 competidores", False),
        ("☐  Calendario de estacionalidad y eventos UY", False),
        ("☐  Benchmark de precios vs competidores", False),
        ("☐  Lista de palabras clave priorizada", False),
        # Fase 3
        ("FASE 3 — AUDITORÍA DE ACTIVOS DIGITALES", True),
        ("☐  Informe de salud del sitio web", False),
        ("☐  Diagnóstico de redes sociales", False),
        ("☐  Tracking y píxeles verificados", False),
        ("☐  Reporte SEO con quick wins", False),
        ("☐  Diagnóstico de email marketing (si aplica)", False),
        # Fase 4
        ("FASE 4 — PRESENCIA DIGITAL", True),
        ("☐  Sitio web optimizado o creado", False),
        ("☐  Perfiles sociales optimizados", False),
        ("☐  Stack analítico configurado", False),
        # Fase 5
        ("FASE 5 — ESTRATEGIA Y PLAN DE MEDIOS", True),
        ("☐  Propuesta comercial y UVP definida", False),
        ("☐  Plan de medios con presupuesto por canal", False),
        ("☐  KPIs y proyecciones documentadas", False),
        # Fase 6
        ("FASE 6 — SETUP TÉCNICO Y CREATIVIDADES", True),
        ("☐  Píxeles y tags funcionando correctamente", False),
        ("☐  Conversiones de Google Ads configuradas", False),
        ("☐  Feed de productos activo (si aplica)", False),
        ("☐  Audiencias creadas y listas", False),
        ("☐  Creatividades diseñadas y aprobadas", False),
        # Fase 7
        ("FASE 7 — PILOTAJE Y TESTEO", True),
        ("☐  Campaña piloto lanzada (7-14 días)", False),
        ("☐  Pruebas A/B en curso", False),
        ("☐  Análisis de resultados completado", False),
        ("☐  Plan de optimización definido", False),
        # Fase 8
        ("FASE 8 — LANZAMIENTO Y ESCALADO", True),
        ("☐  Campañas escaladas gradualmente", False),
        ("☐  Marketing de contenidos activo", False),
        ("☐  Reportes automáticos configurados (n8n)", False),
        ("☐  Ciclo de optimización semanal establecido", False),
    ]

    for item, is_header in checklist_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        if is_header:
            run.bold = True
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_before = Pt(12)
        else:
            run.font.color.rgb = DARK_GRAY
            p.paragraph_format.left_indent = Cm(1.0)

    return doc


# ── Guardar documento ──────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generar documento Word de Metodología Integrada - Global Infinity Marketing"
    )
    parser.add_argument(
        "--output", "-o",
        default=os.path.join(os.path.dirname(os.path.abspath(__file__)), "Metodologia_Integrada_GIM_v2.docx"),
        help="Ruta de salida del documento .docx"
    )
    args = parser.parse_args()

    output_path = args.output
    doc = _build_document()
    doc.save(output_path)
    logger.info(f"Documento generado exitosamente: {output_path}")


if __name__ == "__main__":
    main()
